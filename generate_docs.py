# generate_docs.py
import os
import re
from pathlib import Path
import pathspec
from datetime import datetime
from docx import Document
from docx.shared import Pt

PROJECT_ROOT = Path(__file__).parent
README_FILE = PROJECT_ROOT / "README.md"
DOCS_DIR = PROJECT_ROOT / "documents"

EXTRA_FILES_WITHOUT_EXTENSION = {
    "Dockerfile", "Dockerfile.cpu", "Dockerfile.gpu", "Makefile", ".dockerignore", "LICENSE"
}
INCLUDED_EXTENSIONS = [".py", ".yaml", ".env", ".txt", ".md", ".yml", ".toml", ".sh",'.sql','.api','.webhook','.idoit_webhook','.celery']

EXCLUDED_FILES = {
    ".env", "secrets.py", "config_local.py", "id_rsa", "id_rsa.pub",
    "known_hosts", "docker-compose.override.yml", "token.txt"
}

SYNTAX_HIGHLIGHTING_MAP = {
    ".py": "python",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".sh": "bash",
    ".env": "env",
    ".txt": "",
    ".md": "markdown",
    ".csv": "",
    ".xlsx": "",
    ".json": "json",
    "Dockerfile": "Dockerfile",
    "Makefile": "makefile"
}

# Используем "сырую" строку (r""") или экранируем \
# Также обновим заголовок для лучшего описания
README_HEADER = r"""# 📄 Проект "Ситуационный центр"

## 📝 Описание

Это веб-приложение для мониторинга различных метрик (сеть, логистика, ИТ, ИБ, жалобы) по регионам России. Данные визуализируются на интерактивной карте. Приложение автоматически обновляет данные и отправляет уведомления в Telegram при обнаружении аномалий.

## 🚀 Особенности

- Интерактивная карта с данными по регионам.
- Автоматическая ротация метрик.
- Возможность выбора метрики вручную.
- Уведомления в Telegram о критических значениях.
- Автоматическая генерация документации.
- CI/CD пайплайн для линтинга, тестирования и сборки Docker-образов.

## 🛠️ Установка

1. Клонируйте репозиторий.
2. Создайте виртуальное окружение: `python -m venv venv`
3. Активируйте виртуальное окружение:
   - Windows: `venv\Scripts\activate`
   - Linux/macOS: `source venv/bin/activate`
4. Установите зависимости: `pip install -r requirements.txt`
5. Создайте файл `.env`  и заполните его.
6. Создайте базу данных и сгенерируйте данные: `python generate_data.py`

## ▶️ Запуск



## 🧪 Тестирование

Запустите тесты с помощью `pytest`: `python -m pytest tests/ -v`

## 📁 Структура проекта
""" # Конец строки README_HEADER

def load_gitignore_spec():
    gitignore_path = PROJECT_ROOT / ".gitignore"
    if not gitignore_path.exists():
        return None
    with open(gitignore_path, "r", encoding="utf-8") as f:
        spec = pathspec.PathSpec.from_lines("gitwildmatch", f)
    return spec

def should_include(path, spec):
    rel_path = path.relative_to(PROJECT_ROOT)
    if path.name in EXCLUDED_FILES:
        return False
    return spec is None or not spec.match_file(rel_path)

def generate_tree(start_path, spec):
    tree = []
    for root, dirs, files in os.walk(start_path):
        filtered_dirs = [d for d in dirs if should_include(Path(root) / d, spec)]
        dirs[:] = filtered_dirs
        level = root.replace(str(start_path), "").count(os.sep)
        indent = "│   " * level
        dir_name = os.path.basename(root)
        if level == 0:
            tree.append(f"├── {dir_name}/")
        else:
            tree.append(f"{indent}├── {dir_name}/")
        subindent = "│   " * (level + 1)
        for f in sorted(files):
            full_path = Path(root) / f
            if should_include(full_path, spec):
                if f in EXTRA_FILES_WITHOUT_EXTENSION:
                    tree.append(f"{subindent}├────── {f}")
                elif any(f.endswith(ext) for ext in INCLUDED_EXTENSIONS):
                    tree.append(f"{subindent}├────── {f}")
    return "\n".join(tree).replace("├────── └──", "└──────").replace("├── └──", "└──")

def read_file_content(file_path):
    file_name = file_path.name
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        if file_name in SYNTAX_HIGHLIGHTING_MAP:
            lang = SYNTAX_HIGHLIGHTING_MAP[file_name]
        else:
            lang = SYNTAX_HIGHLIGHTING_MAP.get(file_path.suffix, "")
        return f"```{lang}\n{content}\n```"
    except Exception as e:
        return f"<!-- Ошибка чтения файла: {e} -->\n\n"

def gather_files(start_path, spec):
    file_contents = []
    for root, dirs, files in os.walk(start_path):
        dirs[:] = [d for d in dirs if should_include(Path(root) / d, spec)]
        for f in sorted(files):
            if f == "README.md":
                continue
            full_path = Path(root) / f
            if should_include(full_path, spec):
                if f in EXTRA_FILES_WITHOUT_EXTENSION or any(f.endswith(ext) for ext in INCLUDED_EXTENSIONS):
                    rel_path = full_path.relative_to(PROJECT_ROOT)
                    file_contents.append(f"### 📄 `{rel_path}`\n")
                    file_contents.append(read_file_content(full_path))
    return "\n".join(file_contents)

# --- Новая функция для очистки текста ---
def clean_text_for_xml(text: str) -> str:
    """
    Убираем недопустимые символы для XML 1.0.
    Разрешаем: \t (0x09), \n (0x0A), \r (0x0D) и U+0020..U+D7FF, U+E000..U+FFFD, U+10000..U+10FFFF
    """
    if not isinstance(text, str):
        text = str(text)

    # Удаляем управляющие символы U+0000 - U+001F, кроме таб(9), lf(10), cr(13)
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)
    # Удаляем non-characters U+FFFE, U+FFFF and their high-plane equivalents
    text = re.sub(r'[\uFFFE\uFFFF]', '', text)
    # Также принудительно заменить суррогатные пары, если они появились (условный safe)
    # При кодировке/декодировании с 'utf-8' и errors='ignore' мы убираем невалидные суррогаты
    try:
        text = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
    except Exception:
        # В крайнем случае — убираем всё не-ASCII печатное
        text = ''.join(ch for ch in text if ord(ch) >= 32)

    return text

def build_documentation():
    spec = load_gitignore_spec()
    structure = generate_tree(PROJECT_ROOT, spec)

    # README остаётся лёгким: описание + дерево структуры.
    # Полный дамп исходников в README НЕ пишется — он раздувал файл до нескольких
    # МБ и взрывал git-историю на каждом push. Исходники идут только в .docx-отчёт.
    readme = f"""{README_HEADER}
```
{structure}
```

> Подробнее: [ARCHITECTURE.md](ARCHITECTURE.md) · быстрый старт: [QUICKSTART.md](QUICKSTART.md) · эксплуатация: [docs/operations.md](docs/operations.md)
"""

    # Запись README.md
    with open(README_FILE, "w", encoding="utf-8") as f:
        f.write(readme)
    print("[+] README.md успешно обновлён")

    # Полный листинг кода — отдельный офлайн-артефакт (.docx), не в git.
    code_sections = gather_files(PROJECT_ROOT, spec)
    documentation = f"{readme}\n## 💻 Коды основных модулей\n{code_sections}\n"

    # --- Создание .docx ---
    # Создаем каталог documents, если он не существует
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y_%m_%d_%H_%M")
    project_name = PROJECT_ROOT.name
    docx_filename = f"{project_name}_{timestamp}.docx"
    docx_path = DOCS_DIR / docx_filename

    doc = Document()
    style = doc.styles['Normal']
    font = style.font # type: ignore
    font.name = 'Consolas'
    font.size = Pt(10)

    # Разбиваем всю документацию на строки
    lines = documentation.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        # Очищаем каждую строку перед добавлением
        clean_line = clean_text_for_xml(line)
        
        if clean_line.startswith("```"):
            # Начало блока кода (язык после ``` не используется в .docx)
            code_block = ""
            i += 1
            # Собираем строки кода до закрывающей ```
            while i < len(lines) and not lines[i].startswith("```"):
                # Очищаем и добавляем каждую строку кода
                code_block += clean_text_for_xml(lines[i]) + "\n"
                i += 1
            # Добавляем блок кода в документ
            if code_block.strip():
                doc.add_paragraph(code_block.strip())
                # Простая попытка применить стиль кода (docx не поддерживает подсветку синтаксиса из коробки)
                # Можно рассмотреть использование python-docx-template или других библиотек
                # p_code.style = 'Code' if 'Code' in [s.name for s in doc.styles] else 'Normal'
            i += 1 # Пропускаем закрывающую ```
            continue # Переходим к следующей итерации внешнего цикла
        else:
            # Обычный текст
            if clean_line: # Добавляем только непустые строки
                 p = doc.add_paragraph(clean_line)
                 p.style = doc.styles['Normal'] # type: ignore
        i += 1 # Переход к следующей строке

    try:
        doc.save(docx_path) # type: ignore
        print(f"[+] Документация сохранена как {docx_path}")
    except Exception as e:
        print(f"[!] Ошибка при сохранении .docx файла: {e}")

if __name__ == "__main__":
    build_documentation()
    print("[+] Генерация документации завершена")